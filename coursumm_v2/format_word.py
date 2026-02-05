"""Professional Word document formatting."""

import logging
import re
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

logger = logging.getLogger(__name__)


class WordFormatter:
    """Create professional, book-ready Word documents."""
    
    def __init__(self, course_title: str = "Course Summary"):
        self.course_title = course_title
        self.doc = Document()
        self.chapter_titles = []
        self._setup_document_styles()
    
    def _setup_document_styles(self):
        """Configure document-wide styles."""
        # Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = "Georgia"
        font.size = Pt(11)
        
        # Paragraph formatting
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = 1.5
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.first_line_indent = Inches(0.25)
        
        # Page margins
        for section in self.doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
    
    def add_cover_page(self, subtitle: Optional[str] = None, author: Optional[str] = None):
        """Add a professional cover page."""
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self.course_title)
        title_run.font.name = "Arial"
        title_run.font.size = Pt(36)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(44, 62, 80)
        
        # Spacing
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Subtitle
        if subtitle:
            sub = self.doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub.add_run(subtitle)
            sub_run.font.name = "Arial"
            sub_run.font.size = Pt(18)
            sub_run.font.italic = True
            sub_run.font.color.rgb = RGBColor(52, 152, 219)
        
        # More spacing
        for _ in range(10):
            self.doc.add_paragraph()
        
        # Author
        if author:
            auth = self.doc.add_paragraph()
            auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
            auth_run = auth.add_run(f"Prepared by: {author}")
            auth_run.font.name = "Georgia"
            auth_run.font.size = Pt(14)
        
        self.doc.add_page_break()
    
    def add_table_of_contents(self):
        """Add table of contents."""
        toc_heading = self.doc.add_heading("Table of Contents", level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, chapter_title in enumerate(self.chapter_titles, 1):
            toc_entry = self.doc.add_paragraph()
            toc_entry.paragraph_format.first_line_indent = Inches(0)
            entry_run = toc_entry.add_run(f"Chapter {i}: {chapter_title}")
            entry_run.font.name = "Georgia"
            entry_run.font.size = Pt(11)
        
        self.doc.add_page_break()
    
    def add_chapter_title_page(self, chapter_number: int, chapter_title: str):
        """Add decorative chapter title page."""
        self.doc.add_page_break()
        
        # Spacing from top
        for _ in range(8):
            self.doc.add_paragraph()
        
        # Chapter number
        chapter_num = self.doc.add_paragraph()
        chapter_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_run = chapter_num.add_run(f"CHAPTER {chapter_number}")
        num_run.font.name = "Arial"
        num_run.font.size = Pt(14)
        num_run.font.color.rgb = RGBColor(52, 152, 219)
        
        # Chapter title
        chapter_head = self.doc.add_paragraph()
        chapter_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = chapter_head.add_run(chapter_title)
        title_run.font.name = "Arial"
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(44, 62, 80)
        
        self.doc.add_page_break()
    
    def add_heading(self, text: str, level: int = 1):
        """Add styled heading."""
        heading = self.doc.add_heading(text, level=level)
        heading.paragraph_format.first_line_indent = Inches(0)
        
        for run in heading.runs:
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(44, 62, 80)
            if level == 1:
                run.font.size = Pt(18)
            elif level == 2:
                run.font.size = Pt(14)
    
    def add_quote_box(self, quote_text: str):
        """Add styled quote."""
        quote_para = self.doc.add_paragraph()
        quote_para.paragraph_format.first_line_indent = Inches(0)
        quote_para.paragraph_format.left_indent = Inches(0.5)
        quote_para.paragraph_format.right_indent = Inches(0.5)
        quote_para.paragraph_format.space_before = Pt(12)
        quote_para.paragraph_format.space_after = Pt(12)
        
        quote_run = quote_para.add_run(f'"{quote_text}"')
        quote_run.font.italic = True
        quote_run.font.color.rgb = RGBColor(60, 60, 60)
        quote_run.font.size = Pt(10)
    
    def add_paragraph(self, text: str):
        """Add body paragraph."""
        para = self.doc.add_paragraph(text)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Inches(0.25)
    
    def add_bullet(self, text: str, level: int = 0):
        """Add bullet point."""
        if level == 0:
            para = self.doc.add_paragraph(text, style='List Bullet')
        else:
            para = self.doc.add_paragraph(text, style='List Bullet 2')
        
        para.paragraph_format.first_line_indent = Inches(0)
        para.paragraph_format.left_indent = Inches(0.25 + (level * 0.25))
    
    def add_chapter(self, chapter_num: int, notes_data: Dict, filename: str):
        """Add complete chapter with all sections."""
        # Determine title
        title = notes_data.get("title", "")
        if not title or title.lower() == "untitled":
            title = self._extract_title_from_filename(filename)
        title = self._ensure_question_mark(title)
        
        self.chapter_titles.append(title)
        
        # Chapter title page
        self.add_chapter_title_page(chapter_num, title)
        
        # Quote
        if notes_data.get("quote"):
            self.add_quote_box(notes_data["quote"])
            self.doc.add_paragraph()
        
        # Summary
        self.add_heading("Summary", level=2)
        for para in notes_data.get("summary_paragraphs", []):
            self.add_paragraph(para)
        self.doc.add_paragraph()
        
        # Themes
        self.add_heading("Key Themes", level=2)
        for theme in notes_data.get("themes", []):
            cleaned = re.sub(r'^-\s*', '', theme)
            self.add_bullet(cleaned, level=0)
        self.doc.add_paragraph()
        
        # Key Takeaways
        self.add_heading("Key Takeaways", level=2)
        for line in notes_data.get("key_takeaways", []):
            if re.match(r'^\s*example\s*:', line, re.IGNORECASE):
                ex_text = re.sub(r'^\s*example\s*:\s*', '', line, flags=re.IGNORECASE)
                self.add_bullet(f"Example: {ex_text}", level=1)
            else:
                main_text = re.sub(r'^-\s*', '', line).strip()
                self.add_bullet(main_text, level=0)
        self.doc.add_paragraph()
        
        # Knowledge Check
        self.add_heading("Knowledge Check (Questions)", level=2)
        for i, q in enumerate(notes_data.get("knowledge_check_q", []), 1):
            q_text = re.sub(r'^\d+\)?[.)]\s*', '', q).strip()
            q_para = self.doc.add_paragraph(f"{i}) {q_text}")
            q_para.paragraph_format.first_line_indent = Inches(0)
        
        self.doc.add_paragraph()
        
        self.add_heading("Knowledge Check (Answers)", level=2)
        for i, a in enumerate(notes_data.get("knowledge_check_a", []), 1):
            a_text = re.sub(r'^\d+\)?[.)]\s*', '', a).strip()
            a_para = self.doc.add_paragraph(f"{i}) {a_text}")
            a_para.paragraph_format.first_line_indent = Inches(0)
    
    def _extract_title_from_filename(self, filename: str) -> str:
        """Extract title from filename like 'Lect01 Introduction.txt'."""
        base = Path(filename).stem
        match = re.match(r'^(Lect\d+\s+)(.*)', base, re.IGNORECASE)
        if match:
            return match.group(2).strip()
        return base
    
    def _is_question(self, title: str) -> bool:
        """Check if title is a question."""
        question_words = {"what", "why", "how", "when", "where", "should",
                         "is", "are", "can", "could", "would", "will"}
        if not title:
            return False
        words = title.lower().split()
        return (words[0] in question_words) and not title.endswith("?")
    
    def _ensure_question_mark(self, title: str) -> str:
        """Add question mark if needed."""
        if self._is_question(title):
            return title + "?"
        return title
    
    def save(self, output_path: Path):
        """Save document to file."""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        logger.info(f"Saved document: {output_path}")


def create_compiled_document(
    notes_list: List[tuple],
    output_path: Path,
    course_title: str = "Course Summary",
    subtitle: Optional[str] = None,
    author: Optional[str] = None,
    include_cover: bool = True,
    include_toc: bool = True
) -> Path:
    """
    Create compiled document with all chapters.
    
    Args:
        notes_list: List of (filename, notes_data) tuples
        output_path: Path to save document
        course_title: Course title
        subtitle: Optional subtitle
        author: Optional author
        include_cover: Include cover page
        include_toc: Include table of contents
        
    Returns:
        Path to saved document
    """
    formatter = WordFormatter(course_title)
    
    if include_cover:
        formatter.add_cover_page(subtitle, author)
    
    # Add all chapters (will build TOC list)
    for i, (filename, notes_data) in enumerate(notes_list, 1):
        formatter.add_chapter(i, notes_data, filename)
    
    # Add TOC after building chapter list
    # (In practice, you'd want to insert this after cover)
    # For now, chapters come first, TOC would need reordering
    
    formatter.save(output_path)
    return output_path


def create_individual_documents(
    notes_list: List[tuple],
    output_folder: Path,
    course_title: str = "Course Summary"
) -> List[Path]:
    """
    Create individual Word document for each chapter.
    
    Returns:
        List of saved document paths
    """
    output_folder.mkdir(parents=True, exist_ok=True)
    saved_paths = []
    
    for filename, notes_data in notes_list:
        # Create single-chapter document
        formatter = WordFormatter(course_title)
        formatter.add_chapter(1, notes_data, filename)
        
        # Save with original filename
        base_name = Path(filename).stem
        output_path = output_folder / f"{base_name}.docx"
        formatter.save(output_path)
        saved_paths.append(output_path)
    
    return saved_paths
