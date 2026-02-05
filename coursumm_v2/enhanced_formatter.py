"""
Enhanced Word Formatter for CourseSumm v2 - Phase 2

Professional book-ready formatting with:
- Cover pages
- Table of contents
- Elegant typography
- Decorative elements
- Print-ready margins
- Chapter title pages
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path
from typing import List, Dict
from PIL import Image
import io


class EnhancedWordFormatter:
    """Enhanced Word document formatter for professional book output"""
    
    def __init__(self, course_title, author=None):
        """
        Initialize formatter
        
        Args:
            course_title: Course title
            author: Author/instructor name
        """
        self.course_title = course_title
        self.author = author
        self.doc = Document()
        self._setup_document_margins()
        self._setup_styles()
    
    def _setup_document_margins(self):
        """Set up proper margins for print-ready documents"""
        sections = self.doc.sections
        for section in sections:
            # 1 inch margins all around
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            
            # Letter size (8.5" x 11")
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
    
    def _setup_styles(self):
        """Set up custom styles for the document"""
        styles = self.doc.styles
        
        # Main title style (for cover and chapter titles)
        try:
            title_style = styles['Title']
        except:
            title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        
        title_font = title_style.font
        title_font.name = 'Georgia'
        title_font.size = Pt(36)
        title_font.bold = True
        title_font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-grey
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_before = Pt(72)
        title_style.paragraph_format.space_after = Pt(36)
        
        # Subtitle style
        try:
            subtitle_style = styles['Subtitle']
        except:
            subtitle_style = styles.add_style('Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        
        subtitle_font = subtitle_style.font
        subtitle_font.name = 'Georgia'
        subtitle_font.size = Pt(18)
        subtitle_font.italic = True
        subtitle_font.color.rgb = RGBColor(127, 140, 141)  # Grey
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(24)
        
        # Chapter title style
        try:
            chapter_style = styles['Heading 1']
        except:
            chapter_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        
        chapter_font = chapter_style.font
        chapter_font.name = 'Georgia'
        chapter_font.size = Pt(28)
        chapter_font.bold = True
        chapter_font.color.rgb = RGBColor(52, 73, 94)
        chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chapter_style.paragraph_format.space_before = Pt(144)  # 2 inches
        chapter_style.paragraph_format.space_after = Pt(48)
        chapter_style.paragraph_format.keep_with_next = True
        
        # Section heading style
        try:
            section_style = styles['Heading 2']
        except:
            section_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        
        section_font = section_style.font
        section_font.name = 'Georgia'
        section_font.size = Pt(16)
        section_font.bold = True
        section_font.color.rgb = RGBColor(52, 73, 94)
        section_style.paragraph_format.space_before = Pt(18)
        section_style.paragraph_format.space_after = Pt(12)
        section_style.paragraph_format.keep_with_next = True
        
        # Body text style
        try:
            normal_style = styles['Normal']
        except:
            normal_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
        
        normal_font = normal_style.font
        normal_font.name = 'Georgia'
        normal_font.size = Pt(12)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal_style.paragraph_format.space_after = Pt(12)
        normal_style.paragraph_format.first_line_indent = Inches(0.25)  # First line indent
        
        # Quote style
        try:
            quote_style = styles['Quote']
        except:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
        
        quote_font = quote_style.font
        quote_font.name = 'Georgia'
        quote_font.size = Pt(11)
        quote_font.italic = True
        quote_font.color.rgb = RGBColor(127, 140, 141)
        quote_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.right_indent = Inches(0.5)
        quote_style.paragraph_format.space_before = Pt(18)
        quote_style.paragraph_format.space_after = Pt(18)
        quote_style.paragraph_format.first_line_indent = Inches(0)  # No indent for quotes
    
    def add_cover_page(self, cover_image_path=None, companion_type=None):
        """
        Add a professional cover page
        
        Args:
            cover_image_path: Path to cover image (optional)
            companion_type: Type of companion for subtitle
        """
        # If we have a cover image, insert it
        if cover_image_path and os.path.exists(cover_image_path):
            # Add image centered on page
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            
            # Calculate size to fit page (with margins)
            page_width = Inches(6.5)  # 8.5" - 2" margins
            run.add_picture(cover_image_path, width=page_width)
        else:
            # Text-based cover
            # Add vertical space
            for _ in range(8):
                self.doc.add_paragraph()
            
            # Title
            title = self.doc.add_paragraph(self.course_title, style='Title')
            
            # Companion type
            if companion_type:
                subtitle = self.doc.add_paragraph(companion_type, style='Subtitle')
            
            # Add more space
            for _ in range(4):
                self.doc.add_paragraph()
            
            # Author
            if self.author:
                author_para = self.doc.add_paragraph(self.author)
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                font = author_para.runs[0].font
                font.name = 'Georgia'
                font.size = Pt(14)
        
        # Page break
        self.doc.add_page_break()
    
    def add_decorative_line(self, width_inches=3.0):
        """Add a decorative horizontal line"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run('─' * 30)  # Unicode box drawing character
        font = run.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.color.rgb = RGBColor(149, 165, 166)  # Light grey
    
    def add_table_of_contents(self, chapters: List[Dict]):
        """
        Add a table of contents
        
        Args:
            chapters: List of chapter dictionaries with 'title' keys
        """
        # TOC title
        toc_title = self.doc.add_paragraph('Table of Contents', style='Heading 1')
        
        # Add decorative line
        self.add_decorative_line()
        
        self.doc.add_paragraph()
        
        # Add each chapter
        for i, chapter in enumerate(chapters, 1):
            toc_entry = self.doc.add_paragraph()
            toc_entry.add_run(f"Lecture {i}: ").bold = True
            toc_entry.add_run(chapter.get('title', f'Lecture {i}'))
            
            # Style the entry
            toc_entry.paragraph_format.left_indent = Inches(0.25)
            toc_entry.paragraph_format.space_after = Pt(6)
            font = toc_entry.runs[0].font
            font.name = 'Georgia'
            font.size = Pt(12)
        
        # Page break after TOC
        self.doc.add_page_break()
    
    def add_chapter_title_page(self, lecture_num, title):
        """
        Add a decorative chapter title page
        
        Args:
            lecture_num: Lecture number
            title: Chapter title
        """
        # Add vertical space
        for _ in range(10):
            self.doc.add_paragraph()
        
        # Add decorative element
        self.add_decorative_line()
        
        self.doc.add_paragraph()
        
        # Lecture number
        num_para = self.doc.add_paragraph(f'Lecture {lecture_num}')
        num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_font = num_para.runs[0].font
        num_font.name = 'Georgia'
        num_font.size = Pt(14)
        num_font.color.rgb = RGBColor(127, 140, 141)
        num_font.italic = True
        
        self.doc.add_paragraph()
        
        # Chapter title
        title_para = self.doc.add_paragraph(title, style='Heading 1')
        
        self.doc.add_paragraph()
        
        # Another decorative element
        self.add_decorative_line()
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
    
    def add_quote_box(self, quote_text):
        """
        Add a styled quote box
        
        Args:
            quote_text: Quote text
        """
        # Add the quote with special styling
        quote_para = self.doc.add_paragraph(f'"{quote_text}"', style='Quote')
        
        # Add shading (light background)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F8F9F9')  # Very light grey
        quote_para._element.get_or_add_pPr().append(shading_elm)
        
        # Add border
        pPr = quote_para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # Border size
            border.set(qn('w:space'), '12')  # Space between border and text
            border.set(qn('w:color'), '95A5A6')  # Grey color
            pBdr.append(border)
        pPr.append(pBdr)
    
    def add_section_heading(self, heading_text):
        """Add a section heading (Heading 2)"""
        self.doc.add_paragraph(heading_text, style='Heading 2')
    
    def add_paragraph(self, text, style='Normal'):
        """
        Add a paragraph with proper formatting
        
        Args:
            text: Paragraph text
            style: Style name
        """
        para = self.doc.add_paragraph(text, style=style)
        return para
    
    def add_bullet_points(self, items: List[str], indent_level=0):
        """
        Add bullet points with proper formatting
        
        Args:
            items: List of bullet point texts
            indent_level: Indentation level (0 = main, 1 = sub, etc.)
        """
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25 + (indent_level * 0.25))
            para.paragraph_format.first_line_indent = Inches(-0.25)
            
            # No first-line indent for bullets
            font = para.runs[0].font
            font.name = 'Georgia'
            font.size = Pt(12)
    
    def add_chapter(self, lecture_num, chapter_data: Dict):
        """
        Add a complete chapter with all sections
        
        Args:
            lecture_num: Lecture number
            chapter_data: Dictionary with chapter content
        """
        # Chapter title page
        title = chapter_data.get('title', f'Lecture {lecture_num}')
        self.add_chapter_title_page(lecture_num, title)
        
        # Quote box (if available)
        if 'quote' in chapter_data and chapter_data['quote']:
            self.add_quote_box(chapter_data['quote'])
            self.doc.add_paragraph()
        
        # Summary
        if 'summary' in chapter_data:
            self.add_section_heading('Summary')
            
            # Split summary into paragraphs if it's a single block
            summary_text = chapter_data['summary']
            if isinstance(summary_text, str):
                # Split on double newlines
                paragraphs = summary_text.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        self.add_paragraph(para_text.strip())
            elif isinstance(summary_text, list):
                for para_text in summary_text:
                    self.add_paragraph(para_text)
        
        # Key Themes
        if 'key_themes' in chapter_data:
            self.add_section_heading('Key Themes')
            themes = chapter_data['key_themes']
            
            for i, theme in enumerate(themes, 1):
                if isinstance(theme, dict):
                    theme_title = theme.get('theme', f'Theme {i}')
                    theme_desc = theme.get('description', '')
                    
                    # Theme title as bold paragraph
                    theme_para = self.add_paragraph('')
                    run = theme_para.add_run(f"{i}. {theme_title}")
                    run.bold = True
                    theme_para.paragraph_format.first_line_indent = Inches(0)
                    
                    # Theme description
                    if theme_desc:
                        desc_para = self.add_paragraph(theme_desc)
                        desc_para.paragraph_format.left_indent = Inches(0.25)
                    
                    self.doc.add_paragraph()
                else:
                    self.add_paragraph(f"{i}. {theme}")
        
        # Key Takeaways
        if 'key_takeaways' in chapter_data:
            self.add_section_heading('Key Takeaways')
            takeaways = chapter_data['key_takeaways']
            
            if isinstance(takeaways, list):
                for takeaway in takeaways:
                    if isinstance(takeaway, dict):
                        point = takeaway.get('point', '')
                        example = takeaway.get('example', '')
                        
                        self.add_bullet_points([point])
                        if example:
                            ex_para = self.doc.add_paragraph(f"Example: {example}")
                            ex_para.paragraph_format.left_indent = Inches(0.5)
                            ex_para.paragraph_format.first_line_indent = Inches(0)
                            font = ex_para.runs[0].font
                            font.italic = True
                            font.size = Pt(11)
                    else:
                        self.add_bullet_points([takeaway])
        
        # Knowledge Check
        if 'knowledge_check' in chapter_data:
            self.add_section_heading('Knowledge Check')
            qa_pairs = chapter_data['knowledge_check']
            
            for i, qa in enumerate(qa_pairs, 1):
                if isinstance(qa, dict):
                    question = qa.get('question', '')
                    answer = qa.get('answer', '')
                    
                    # Question
                    q_para = self.add_paragraph('')
                    q_para.paragraph_format.first_line_indent = Inches(0)
                    run = q_para.add_run(f"Q{i}: {question}")
                    run.bold = True
                    
                    # Answer
                    a_para = self.add_paragraph(f"A: {answer}")
                    a_para.paragraph_format.left_indent = Inches(0.25)
                    a_para.paragraph_format.first_line_indent = Inches(0)
                    
                    self.doc.add_paragraph()
        
        # Page break after chapter
        self.doc.add_page_break()
    
    def save(self, output_path):
        """
        Save the document
        
        Args:
            output_path: Where to save the document
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        return output_path


def format_enhanced_document(course_title, chapters: List[Dict], 
                            output_path, author=None,
                            cover_image_path=None, companion_type=None):
    """
    Create a complete enhanced Word document
    
    Args:
        course_title: Course title
        chapters: List of chapter data dictionaries
        output_path: Output file path
        author: Author name
        cover_image_path: Path to cover image
        companion_type: Type of companion for subtitle
        
    Returns:
        Path to saved document
    """
    formatter = EnhancedWordFormatter(course_title, author)
    
    # Add cover page
    formatter.add_cover_page(cover_image_path, companion_type)
    
    # Add table of contents
    formatter.add_table_of_contents(chapters)
    
    # Add all chapters
    for i, chapter in enumerate(chapters, 1):
        formatter.add_chapter(i, chapter)
    
    # Save
    return formatter.save(output_path)
