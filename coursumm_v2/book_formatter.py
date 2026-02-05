"""
Professional Book Formatter for CourseSumm v2

Creates sale-ready Word documents with:
- Professional title page
- Copyright page with legal disclaimer
- Table of contents
- Page numbers and headers
- About the Author
- Proper typography
"""

import os
from datetime import datetime
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ProfessionalBookFormatter:
    """Creates professional, sale-ready book documents"""
    
    # Standard book metadata
    SERIES_NAME = "Great Course Companions"
    PUBLISHER = "Independent Publishing"
    YEAR = datetime.now().year
    
    # Legal disclaimer
    DISCLAIMER = """This companion guide is an independent educational resource created to enhance your learning experience. It is not affiliated with, endorsed by, or connected to The Great Courses, The Teaching Company, or any course instructor.

All original content, analysis, discussion questions, and synthesis in this guide are the intellectual property of the author. This guide is intended as a transformative educational work that provides original commentary, analysis, and study materials.

No part of this publication may be reproduced, distributed, or transmitted in any form without prior written permission of the publisher, except for brief quotations in reviews and educational contexts."""

    def __init__(self, 
                 course_title: str,
                 author: str = "Seo-Yun Kim",
                 companion_type: str = "Lecture Companion",
                 subtitle: str = None):
        self.course_title = course_title
        self.author = author
        self.companion_type = companion_type
        self.subtitle = subtitle or f"A {companion_type} Guide"
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Set up professional typography styles"""
        styles = self.doc.styles
        
        # Normal text - Garamond-like (use Georgia as fallback)
        normal = styles['Normal']
        normal.font.name = 'Georgia'
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal.paragraph_format.space_after = Pt(6)
        
        # Title style - use existing Heading style or create
        try:
            title_style = styles['Title']
            title_style.font.name = 'Arial'
            title_style.font.size = Pt(28)
            title_style.font.bold = True
        except KeyError:
            pass
        
        # Subtitle style - use existing or skip
        try:
            sub_style = styles['Subtitle']
            sub_style.font.name = 'Arial'
            sub_style.font.size = Pt(16)
            sub_style.font.italic = True
        except KeyError:
            pass
        
        # Chapter title
        h1 = styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.page_break_before = True
        
        # Section heading
        h2 = styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(6)
        
        # Subsection
        h3 = styles['Heading 3']
        h3.font.name = 'Arial'
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)
        
        # Quote style - use Quote if available
        try:
            quote_style = styles['Quote']
            quote_style.font.name = 'Georgia'
            quote_style.font.size = Pt(11)
            quote_style.font.italic = True
        except KeyError:
            pass
    
    def _add_page_break(self):
        """Add a page break"""
        self.doc.add_page_break()
    
    def _add_title_page(self, cover_image_path: str = None):
        """Add professional title page"""
        # Series name at top
        series = self.doc.add_paragraph(self.SERIES_NAME)
        series.alignment = WD_ALIGN_PARAGRAPH.CENTER
        series.runs[0].font.size = Pt(12)
        series.runs[0].font.italic = True
        series.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        
        # Spacer
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Cover image if provided
        if cover_image_path and os.path.exists(cover_image_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(cover_image_path, width=Inches(4))
            self.doc.add_paragraph()
        
        # Main title
        title = self.doc.add_paragraph(self.course_title)
        title.style = 'Title'
        
        # Subtitle
        subtitle = self.doc.add_paragraph(self.subtitle)
        subtitle.style = 'Subtitle'
        
        # Companion type badge
        badge = self.doc.add_paragraph(f"— {self.companion_type} —")
        badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
        badge.runs[0].font.size = Pt(12)
        badge.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        
        # Spacer
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Author
        author_p = self.doc.add_paragraph(f"By {self.author}")
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_p.runs[0].font.size = Pt(14)
        
        self._add_page_break()
    
    def _add_copyright_page(self):
        """Add copyright and legal disclaimer page"""
        # Copyright notice
        copyright_text = f"""Copyright © {self.YEAR} {self.author}
All rights reserved.

Published by {self.PUBLISHER}

First Edition: {datetime.now().strftime('%B %Y')}

ISBN: [Pending]

"""
        p = self.doc.add_paragraph(copyright_text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10)
        
        # Legal disclaimer
        self.doc.add_paragraph()
        disclaimer_title = self.doc.add_paragraph("IMPORTANT NOTICE")
        disclaimer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer_title.runs[0].font.bold = True
        disclaimer_title.runs[0].font.size = Pt(10)
        
        disclaimer = self.doc.add_paragraph(self.DISCLAIMER)
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        disclaimer.runs[0].font.size = Pt(9)
        
        self._add_page_break()
    
    def _add_toc_placeholder(self):
        """Add table of contents (placeholder - Word will generate)"""
        toc_title = self.doc.add_paragraph("Table of Contents")
        toc_title.style = 'Heading 1'
        toc_title.paragraph_format.page_break_before = False
        
        # Instruction for Word
        instruction = self.doc.add_paragraph(
            "[To generate: In Word, go to References → Table of Contents → Automatic Table 1]"
        )
        instruction.runs[0].font.italic = True
        instruction.runs[0].font.size = Pt(10)
        instruction.runs[0].font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
        
        # Manual TOC entries will be added by content methods
        self.doc.add_paragraph()
        self._add_page_break()
    
    def _add_introduction(self, intro_text: str = None):
        """Add introduction section"""
        intro_title = self.doc.add_paragraph("Introduction")
        intro_title.style = 'Heading 1'
        intro_title.paragraph_format.page_break_before = False
        
        if intro_text:
            self.doc.add_paragraph(intro_text)
        else:
            default_intro = f"""Welcome to this companion guide for "{self.course_title}."

This guide is designed to enhance your learning experience by providing structured summaries, key concepts, discussion questions, and deeper analysis of the course material. Whether you're reviewing after lectures or preparing for discussions, this companion will help you engage more deeply with the philosophical ideas presented.

Each section includes:
• **Key Points** — The main arguments and ideas from each lecture
• **Core Concepts** — Definitions and explanations of important terms
• **Discussion Questions** — Prompts to stimulate critical thinking
• **Practical Applications** — Ways to apply these ideas to everyday life

Use this guide alongside your course materials to maximize your understanding and retention of these profound philosophical questions.

Happy learning!

{self.author}"""
            self.doc.add_paragraph(default_intro)
        
        self._add_page_break()
    
    def _add_about_author(self):
        """Add About the Author page"""
        about_title = self.doc.add_paragraph("About the Author")
        about_title.style = 'Heading 1'
        
        about_text = f"""{self.author} is an independent educational writer and course companion author specializing in making complex academic content accessible to lifelong learners.

With a passion for philosophy, science, and the humanities, {self.author.split()[0]} creates study guides that help students and curious minds engage more deeply with great courses and lectures from top professors.

The {self.SERIES_NAME} series provides thoughtful analysis, discussion questions, and synthesis to transform passive viewing into active learning.

Connect with the author:
• Email: [contact email]
• Website: [website]

Other titles in the {self.SERIES_NAME} series:
[More titles coming soon]"""
        
        self.doc.add_paragraph(about_text)
    
    def _add_headers_footers(self):
        """Add running headers and page numbers"""
        for section in self.doc.sections:
            # Header
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = f"{self.course_title}  |  {self.companion_type}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.runs[0].font.size = Pt(9) if header_para.runs else None
            header_para.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D) if header_para.runs else None
            
            # Footer with page number
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add page number field
            run = footer_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            
            run.font.size = Pt(10)
    
    def add_lecture_chapter(self, lecture_num: int, lecture_data: Dict[str, Any]):
        """Add a lecture as a chapter with proper formatting"""
        title = lecture_data.get('title', f'Lecture {lecture_num}')
        
        # Chapter heading - lecture number on one line, title on next
        chapter_title = self.doc.add_paragraph()
        chapter_title.style = 'Heading 1'
        run1 = chapter_title.add_run(f"Lecture {lecture_num}")
        run1.bold = True
        chapter_title.add_run("\n")
        run2 = chapter_title.add_run(title)
        
        # Introduction
        if 'introduction' in lecture_data:
            intro = self.doc.add_paragraph(lecture_data['introduction'])
            intro.style = 'Quote'
        
        # Main Points
        if 'main_points' in lecture_data:
            self.doc.add_paragraph("Key Points", style='Heading 2')
            for i, point in enumerate(lecture_data['main_points'], 1):
                if isinstance(point, dict):
                    point_title = self.doc.add_paragraph(f"{i}. {point.get('point', '')}")
                    point_title.runs[0].font.bold = True
                    if point.get('explanation'):
                        self.doc.add_paragraph(point['explanation'])
                else:
                    self.doc.add_paragraph(f"{i}. {point}")
        
        # Key Concepts
        if 'key_concepts' in lecture_data:
            self.doc.add_paragraph("Core Concepts", style='Heading 2')
            for concept in lecture_data['key_concepts']:
                if isinstance(concept, dict):
                    concept_p = self.doc.add_paragraph()
                    concept_name = concept_p.add_run(f"• {concept.get('concept', '')}: ")
                    concept_name.bold = True
                    concept_p.add_run(concept.get('definition', ''))
                else:
                    self.doc.add_paragraph(f"• {concept}")
        
        # Practical Applications
        if 'practical_applications' in lecture_data:
            self.doc.add_paragraph("Practical Applications", style='Heading 2')
            for app in lecture_data['practical_applications']:
                if isinstance(app, dict):
                    app_p = self.doc.add_paragraph()
                    app_name = app_p.add_run(f"• {app.get('application', '')}: ")
                    app_name.bold = True
                    app_p.add_run(app.get('explanation', ''))
                else:
                    self.doc.add_paragraph(f"• {app}")
        
        # Discussion Questions
        if 'discussion_questions' in lecture_data:
            self.doc.add_paragraph("Discussion Questions", style='Heading 2')
            for i, q in enumerate(lecture_data['discussion_questions'], 1):
                self.doc.add_paragraph(f"{i}. {q}")
        
        # Further Exploration
        if 'further_exploration' in lecture_data:
            self.doc.add_paragraph("Further Exploration", style='Heading 2')
            for item in lecture_data['further_exploration']:
                self.doc.add_paragraph(f"• {item}")
    
    def add_synthesis_chapter(self, synthesis_data: Dict[str, Any]):
        """Add synthesis/going deeper chapter"""
        # Introduction
        if 'introduction' in synthesis_data:
            intro_title = self.doc.add_paragraph("Course Overview")
            intro_title.style = 'Heading 1'
            intro_title.paragraph_format.page_break_before = False
            self.doc.add_paragraph(synthesis_data['introduction'])
        
        # Major Themes
        if 'major_themes' in synthesis_data:
            themes_title = self.doc.add_paragraph("Major Themes")
            themes_title.style = 'Heading 1'
            
            for theme in synthesis_data['major_themes']:
                if isinstance(theme, dict):
                    theme_name = self.doc.add_paragraph(theme.get('theme', ''))
                    theme_name.style = 'Heading 2'
                    
                    if theme.get('description'):
                        self.doc.add_paragraph(theme['description'])
                    
                    if theme.get('key_lectures'):
                        kl = self.doc.add_paragraph()
                        kl.add_run("Key Lectures: ").bold = True
                        kl.add_run(theme['key_lectures'])
                    
                    if theme.get('connections'):
                        conn = self.doc.add_paragraph()
                        conn.add_run("Connections: ").bold = True
                        conn.add_run(theme['connections'])
        
        # Deeper Questions
        if 'deeper_questions' in synthesis_data:
            dq_title = self.doc.add_paragraph("Deeper Questions")
            dq_title.style = 'Heading 1'
            
            for q in synthesis_data['deeper_questions']:
                if isinstance(q, dict):
                    q_title = self.doc.add_paragraph(q.get('question', ''))
                    q_title.style = 'Heading 2'
                    
                    if q.get('exploration'):
                        self.doc.add_paragraph(q['exploration'])
                    
                    if q.get('further_reading'):
                        fr = self.doc.add_paragraph()
                        fr.add_run("Further Reading: ").bold = True
                        fr.add_run(q['further_reading'])
        
        # Intellectual Connections
        if 'intellectual_connections' in synthesis_data:
            ic = synthesis_data['intellectual_connections']
            ic_title = self.doc.add_paragraph("Intellectual Connections")
            ic_title.style = 'Heading 1'
            
            if ic.get('related_fields'):
                rf = self.doc.add_paragraph()
                rf.add_run("Related Fields: ").bold = True
                rf.add_run(", ".join(ic['related_fields']))
            
            if ic.get('key_thinkers'):
                kt = self.doc.add_paragraph()
                kt.add_run("Key Thinkers: ").bold = True
                kt.add_run(", ".join(ic['key_thinkers']))
            
            if ic.get('modern_relevance'):
                self.doc.add_paragraph()
                self.doc.add_paragraph(ic['modern_relevance'])
        
        # Final Synthesis
        if 'synthesis' in synthesis_data:
            syn_title = self.doc.add_paragraph("Synthesis")
            syn_title.style = 'Heading 1'
            self.doc.add_paragraph(synthesis_data['synthesis'])
    
    def build_v1_book(self, lectures: List[Dict], cover_path: str = None) -> Document:
        """Build complete V1 (Lecture Companion) book"""
        self.companion_type = "Lecture Companion"
        self.subtitle = "Your Guide to Each Lecture"
        
        self._add_title_page(cover_path)
        self._add_copyright_page()
        self._add_toc_placeholder()
        self._add_introduction()
        
        for i, lecture in enumerate(lectures, 1):
            self.add_lecture_chapter(i, lecture)
        
        self._add_about_author()
        self._add_headers_footers()
        
        return self.doc
    
    def build_v2_book(self, synthesis_data: Dict, cover_path: str = None) -> Document:
        """Build complete V2 (Going Deeper) book"""
        self.companion_type = "Going Deeper"
        self.subtitle = "Themes, Connections & Questions"
        
        self._add_title_page(cover_path)
        self._add_copyright_page()
        self._add_toc_placeholder()
        self.add_synthesis_chapter(synthesis_data)
        self._add_about_author()
        self._add_headers_footers()
        
        return self.doc
    
    def build_v3_book(self, v3_data: Dict, cover_path: str = None) -> Document:
        """Build complete V3 (Complete Companion) book"""
        self.companion_type = "Complete Companion"
        self.subtitle = "The Comprehensive Guide"
        
        self._add_title_page(cover_path)
        self._add_copyright_page()
        self._add_toc_placeholder()
        
        # Part 1: Lecture-by-lecture
        if 'lectures' in v3_data:
            part1 = self.doc.add_paragraph("Part I: Lecture Companions")
            part1.style = 'Heading 1'
            part1.runs[0].font.size = Pt(24)
            
            for i, lecture in enumerate(v3_data['lectures'], 1):
                self.add_lecture_chapter(i, lecture)
        
        # Part 2: Synthesis
        if 'synthesis' in v3_data:
            part2 = self.doc.add_paragraph("Part II: Going Deeper")
            part2.style = 'Heading 1'
            part2.runs[0].font.size = Pt(24)
            
            self.add_synthesis_chapter(v3_data['synthesis'])
        
        self._add_about_author()
        self._add_headers_footers()
        
        return self.doc
    
    def save(self, output_path: str):
        """Save the document"""
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        return output_path


def create_professional_book(course_title: str, 
                            content_data: Dict,
                            output_path: str,
                            book_type: str = "v1",
                            author: str = "Seo-Yun Kim",
                            cover_path: str = None) -> str:
    """
    Convenience function to create a professional book
    
    Args:
        course_title: Title of the course
        content_data: Dictionary with content (lectures, synthesis, etc.)
        output_path: Where to save the .docx
        book_type: "v1", "v2", or "v3"
        author: Author name
        cover_path: Path to cover image
    
    Returns:
        Path to saved file
    """
    formatter = ProfessionalBookFormatter(
        course_title=course_title,
        author=author
    )
    
    if book_type == "v1":
        formatter.build_v1_book(content_data.get('lectures', []), cover_path)
    elif book_type == "v2":
        formatter.build_v2_book(content_data, cover_path)
    elif book_type == "v3":
        formatter.build_v3_book(content_data, cover_path)
    else:
        raise ValueError(f"Unknown book type: {book_type}")
    
    formatter.save(output_path)
    return output_path
